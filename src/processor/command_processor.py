from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import sys
from typing import Dict, List, Optional, Union, TypedDict

class StyleSpec(TypedDict, total=False):
    font: str
    size: int
    bold: bool
    italic: bool
    underline: bool
    color: str
    style: str
    spacing: Dict[str, int]
    alignment: str

class ContentSpec(TypedDict):
    text: str
    style: Dict[str, Union[bool, StyleSpec]]

class LocationSpec(TypedDict):
    type: str
    value: Optional[str]
    number: Optional[Union[str, int]]
    position: str
    matchText: Optional[str]
    section: Optional[Union[str, int]]
    paragraphNumber: Optional[int]
    sentenceNumber: Optional[int]
    matchLevel: Optional[bool]

class DocumentCommand(TypedDict):
    documentId: str
    action: str
    location: LocationSpec
    content: Optional[ContentSpec]
    validation: Optional[Dict]

class CommandProcessor:
    def __init__(self, doc_path: str):
        self.doc = Document(doc_path)
        self.doc_path = doc_path

    def find_heading(self, heading_text: str) -> Optional[int]:
        """Find paragraph index of a heading"""
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name.startswith('Heading') and para.text.strip() == heading_text.strip():
                return i
        return None

    def find_section(self, section_num: Union[str, int]) -> Optional[int]:
        """Find paragraph index of a section start"""
        section_str = str(section_num)
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.startswith(section_str + '.'):
                return i
        return None

    def find_sentence(self, location: LocationSpec) -> Optional[int]:
        """Find paragraph and position of a sentence"""
        start_para = 0
        if location.get('section'):
            start_para = self.find_section(location['section']) or 0

        for i in range(start_para, len(self.doc.paragraphs)):
            para = self.doc.paragraphs[i]
            if location.get('matchText') and location['matchText'] in para.text:
                return i
            elif location.get('sentenceNumber'):
                # Simple sentence splitting - could be improved
                sentences = para.text.split('. ')
                if len(sentences) >= location['sentenceNumber']:
                    return i
        return None

    def apply_style(self, paragraph, style_spec: StyleSpec):
        """Apply style specifications to a paragraph"""
        if style_spec.get('style'):
            paragraph.style = style_spec['style']
        
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        
        if style_spec.get('font'):
            run.font.name = style_spec['font']
        if style_spec.get('size'):
            run.font.size = Pt(style_spec['size'])
        if style_spec.get('bold') is not None:
            run.font.bold = style_spec['bold']
        if style_spec.get('italic') is not None:
            run.font.italic = style_spec['italic']
        if style_spec.get('underline') is not None:
            run.font.underline = style_spec['underline']
        if style_spec.get('color'):
            color = style_spec['color'].lstrip('#')
            run.font.color.rgb = RGBColor(
                int(color[:2], 16),
                int(color[2:4], 16),
                int(color[4:], 16)
            )
        
        if style_spec.get('spacing'):
            spacing = style_spec['spacing']
            if 'before' in spacing:
                paragraph.paragraph_format.space_before = Pt(spacing['before'])
            if 'after' in spacing:
                paragraph.paragraph_format.space_after = Pt(spacing['after'])
            if 'line' in spacing:
                paragraph.paragraph_format.line_spacing = spacing['line']
        
        if style_spec.get('alignment'):
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            paragraph.alignment = align_map[style_spec['alignment']]

    def validate_conditions(self, conditions: Dict[str, List[str]], pre: bool = True) -> bool:
        """Validate pre/post conditions"""
        doc_text = ' '.join(p.text for p in self.doc.paragraphs)
        
        if pre:
            must_exist = conditions.get('mustExist', [])
            must_not_exist = conditions.get('mustNotExist', [])
        else:
            must_exist = conditions.get('shouldExist', [])
            must_not_exist = conditions.get('shouldNotExist', [])

        for text in must_exist:
            if text not in doc_text:
                return False
        
        for text in must_not_exist:
            if text in doc_text:
                return False
        
        return True

    def process_command(self, command: DocumentCommand) -> bool:
        """Process a document command"""
        # Validate pre-conditions
        if command.get('validation') and command['validation'].get('preConditions'):
            if not self.validate_conditions(command['validation']['preConditions'], pre=True):
                raise ValueError("Pre-conditions not met")

        location = command['location']
        target_idx = None

        # Find target location
        if location['type'] == 'heading':
            target_idx = self.find_heading(location['value'])
        elif location['type'] == 'section':
            target_idx = self.find_section(location['number'])
        elif location['type'] in ['sentence', 'paragraph']:
            target_idx = self.find_sentence(location)

        if target_idx is None:
            raise ValueError(f"Could not find target location: {location}")

        # Execute command
        if command['action'] == 'delete':
            self.doc.paragraphs[target_idx]._element.getparent().remove(
                self.doc.paragraphs[target_idx]._element
            )
        else:  # insert or modify
            content = command['content']
            if not content:
                raise ValueError("Content required for insert/modify actions")

            if command['action'] == 'insert':
                if location['position'] == 'before':
                    target_idx -= 1
                new_para = self.doc.add_paragraph()
                if target_idx < len(self.doc.paragraphs) - 1:
                    target_para = self.doc.paragraphs[target_idx + 1]
                    new_para._p.addnext(target_para._p)
            else:  # modify
                new_para = self.doc.paragraphs[target_idx]

            new_para.text = content['text']
            if content.get('style'):
                if content['style'].get('matchSource') and target_idx >= 0:
                    # Copy style from target
                    new_para.style = self.doc.paragraphs[target_idx].style
                if content['style'].get('specific'):
                    self.apply_style(new_para, content['style']['specific'])

        # Save changes
        self.doc.save(self.doc_path)

        # Validate post-conditions
        if command.get('validation') and command['validation'].get('postConditions'):
            if not self.validate_conditions(command['validation']['postConditions'], pre=False):
                raise ValueError("Post-conditions not met")

        return True

def main():
    if len(sys.argv) != 3:
        print(json.dumps({"status": "error", "message": "Invalid arguments"}))
        sys.exit(1)

    doc_path = sys.argv[1]
    command_json = sys.argv[2]

    try:
        command = json.loads(command_json)
        processor = CommandProcessor(doc_path)
        processor.process_command(command)
        print(json.dumps({"status": "success"}))
    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}))

if __name__ == "__main__":
    main() 