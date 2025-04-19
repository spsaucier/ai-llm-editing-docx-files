from docx import Document
from typing import Dict, Tuple, Optional
import json
import sys

class DocumentProcessor:
    def __init__(self, doc_path: str):
        self.doc = Document(doc_path)
        self._analyze_styles()

    def _analyze_styles(self) -> None:
        """Analyze document styles for later matching"""
        self.styles = {
            p.style.name: {
                "bold": p.style.font.bold,
                "underline": p.style.font.underline
            } for p in self.doc.paragraphs if p.style and p.style.font
        }

    def _find_section(self, section_id: str) -> Tuple[int, Optional[str]]:
        """Find a section by its ID (e.g., '4.1')"""
        for i, p in enumerate(self.doc.paragraphs):
            if p.text.startswith(section_id):
                style_name = p.style.name if p.style else None
                return i, style_name
        return -1, None

    def process_document(
        self,
        clause: str,
        target_section: str,
        formatting: Dict
    ) -> bool:
        """Insert a clause after the specified section"""
        pos, style_name = self._find_section(target_section)
        if pos == -1:
            return False

        # Find the end of the section
        next_pos = pos + 1
        while (next_pos < len(self.doc.paragraphs) and 
               not any(self.doc.paragraphs[next_pos].text.startswith(str(i)) 
                      for i in range(10))):
            next_pos += 1

        # Insert the new clause
        p = self.doc.paragraphs[pos]
        new_p = self.doc.add_paragraph()
        self.doc._body._body.insert(next_pos, new_p._p)
        new_p.text = clause

        # Apply formatting
        if style_name:
            new_p.style = self.doc.styles[style_name]
        if formatting.get("bold"):
            new_p.runs[0].bold = True
        if formatting.get("underline"):
            new_p.runs[0].underline = True

        return True

def main():
    """CLI interface for document processing"""
    if len(sys.argv) != 5:
        print("Usage: processor.py <doc_path> <clause> <target_section> <formatting_json>")
        sys.exit(1)

    doc_path = sys.argv[1]
    clause = sys.argv[2]
    target_section = sys.argv[3]
    formatting = json.loads(sys.argv[4])

    processor = DocumentProcessor(doc_path)
    success = processor.process_document(clause, target_section, formatting)
    
    if success:
        processor.doc.save(doc_path)
        print(json.dumps({"status": "success"}))
    else:
        print(json.dumps({"status": "error", "message": "Section not found"}))

if __name__ == "__main__":
    main() 