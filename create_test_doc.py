from docx import Document

doc = Document()
doc.add_paragraph('section1. Test Section')
doc.add_paragraph('This is a test section content.')
doc.add_paragraph('section2. Another Section')
doc.add_paragraph('This is another section content.')

doc.save('src/processor/test.docx') 